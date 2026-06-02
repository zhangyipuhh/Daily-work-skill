#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
generate_report.py - 最终 DOCX 报告生成脚本（V2.1）

读取：
  - 当前周的 review_results.json
  - 当前周的 文档审查结果_*.xlsx
  - 趋势分析_*.md（可选）
  - 漏检名单 JSON（可选）
  - baseline.db（SQLite，第七章用）

输出：AI 辅助编程报告_<week>.docx

章节（7 + 1 模型，严格匹配 references/report_format_spec.md）：
  0. 标题 + 副标题
  1. 第一章 概述
  2. 第二章 汇总统计
  3. 第三章 评审结果汇总（Top/Bottom 开发者）
  4. 第四章 改进建议汇总
  5. 第五章 不足与改进建议
     5.x 本周未提交者（漏检名单，嵌入子章节）
  6. 第六章 整体趋势变化分析
  7. 总结（独立章节）
  8. 第七章 全量比对（最近 4 次，从 baseline.db 读取）

依赖：openpyxl, python-docx（sqlite3 为 Python 内置）
"""
import argparse
import json
import sqlite3
import sys
from pathlib import Path
from datetime import datetime
from collections import Counter, defaultdict

try:
    from openpyxl import load_workbook
except ImportError:
    print("[ERROR] 缺少 openpyxl，请先 pip install openpyxl", file=sys.stderr)
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("[ERROR] 缺少 python-docx，请先 pip install python-docx", file=sys.stderr)
    sys.exit(1)


BLACK = RGBColor(0, 0, 0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_FILL = "4472C4"
ALT_FILL = "F2F2F2"
TREND_UP = "↑ 提升"
TREND_DOWN = "↓ 下降"
TREND_FLAT = "— 持平"
TREND_NEW = "NEW 新人"


def set_run_fonts(run, ascii_font="Arial", east_asia_font="微软雅黑"):
    """统一设置 run 字体（ASCII + 东亚双字体）"""
    run.font.name = ascii_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), east_asia_font)
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)


def set_cell_text(cell, text, *, bold=False, color=BLACK, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text) if text is not None else "")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    set_run_fonts(run)
    return cell


def set_cell_shading(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    run.font.color.rgb = BLACK
    run.bold = True
    set_run_fonts(run)
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return h


def add_paragraph(doc, text, *, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    run.bold = bold
    set_run_fonts(run)
    return p


def add_bullet(doc, text, *, size=11):
    """列表项（左缩进 + • 前缀）"""
    p = doc.add_paragraph(style="List Paragraph")
    run = p.add_run("• " + text)
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    set_run_fonts(run)
    return p


def add_table(doc, headers, rows, *, header_fill=HEADER_FILL):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_text(cell, h, bold=True, color=WHITE)
        set_cell_shading(cell, header_fill)
    for r_idx, row in enumerate(rows, 1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            display = "—" if val is None or val == "" else val
            set_cell_text(cell, display)
            if r_idx % 2 == 0:
                set_cell_shading(cell, ALT_FILL)
    return table


def safe_get(d, *keys, default=None):
    cur = d
    for k in keys:
        if not isinstance(cur, dict):
            return default
        cur = cur.get(k)
        if cur is None:
            return default
    return cur


def load_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


def load_excel_rows(path: Path):
    if not path.exists():
        return [], []
    wb = load_workbook(path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))
    if not rows:
        return [], []
    headers = [str(c) if c is not None else "" for c in rows[0]]
    data = []
    for r in rows[1:]:
        data.append([("" if v is None else v) for v in r])
    return headers, data


def collect_overview(json_path: Path, week_label: str) -> dict:
    data = load_json(json_path)
    developers = []
    for item in data:
        if not isinstance(item, dict):
            continue
        d = item.get("data", {})
        if not isinstance(d, dict):
            continue
        developers.append(d)
    return {"week": week_label, "count": len(developers), "developers": developers}


# ========================================================================
# 章节函数（按 references/report_format_spec.md 定义的 7+1 模型）
# ========================================================================

def chapter_summary(doc, overview, json_path):
    add_heading(doc, "第一章 概述", level=1)
    add_paragraph(
        doc,
        f"本报告基于周次 {overview['week']} 的 AI 辅助编程提交材料评审。",
    )
    add_paragraph(
        doc,
        f"评审覆盖开发者 {overview['count']} 人，数据源: {json_path.name}。",
    )
    add_paragraph(
        doc,
        "评审聚焦于：文档质量、文档与代码/任务的一致性、AI 编程采纳率、反模式检测。",
    )


def chapter_stats(doc, overview):
    devs = overview["developers"]
    add_heading(doc, "第二章 汇总统计", level=1)

    overall = [safe_get(d, "overall_score") for d in devs if isinstance(safe_get(d, "overall_score"), (int, float))]
    docq = [safe_get(d, "document_quality", "overall_score") for d in devs if isinstance(safe_get(d, "document_quality", "overall_score"), (int, float))]
    aiad = [safe_get(d, "ai_adoption_rate", "adoption_rate") for d in devs if isinstance(safe_get(d, "ai_adoption_rate", "adoption_rate"), (int, float))]
    consis = [safe_get(d, "ai_adoption_rate", "doc_code_consistency", "consistency_score") for d in devs if isinstance(safe_get(d, "ai_adoption_rate", "doc_code_consistency", "consistency_score"), (int, float))]
    anti = [safe_get(d, "anti_patterns", "has_issues") for d in devs]

    def avg(vs):
        return sum(vs) / len(vs) if vs else 0

    def pct(n, t):
        return f"{(n / t * 100):.1f}%" if t else "0.0%"

    valid = len(overall)
    no_score = len(devs) - valid
    anti_count = sum(1 for x in anti if x is True)

    rows = [
        ["评审覆盖开发者", f"{len(devs)} 人"],
        ["有效评分人数", f"{valid} 人"],
        ["无评分/数据不足", f"{no_score} 人 ({pct(no_score, len(devs))})"],
        ["平均综合评分", f"{avg(overall):.2f}"],
        ["最高分", f"{max(overall):.2f}" if overall else "—"],
        ["最低分", f"{min(overall):.2f}" if overall else "—"],
        ["平均文档质量", f"{avg(docq):.2f}"],
        ["平均 AI 采纳率", f"{avg(aiad):.2%}" if aiad else "—"],
        ["平均文档代码一致性", f"{avg(consis):.2%}" if consis else "—"],
        ["反模式命中人数", f"{anti_count} 人"],
    ]
    add_table(doc, ["指标", "数值"], rows)


def chapter_top_bottom(doc, overview):
    add_heading(doc, "第三章 评审结果汇总", level=1)
    devs = overview["developers"]
    scored = []
    for d in devs:
        s = safe_get(d, "overall_score")
        if isinstance(s, (int, float)):
            scored.append({
                "name": safe_get(d, "name", default="未知"),
                "score": s,
                "ai": safe_get(d, "ai_adoption_rate", "adoption_rate"),
                "priority": safe_get(d, "improvement_priority") or "—",
            })
    scored.sort(key=lambda x: x["score"], reverse=True)

    if scored:
        add_paragraph(doc, "### 3.1 Top 10 高分开发者", bold=True, size=12)
        rows = []
        for i, item in enumerate(scored[:10], 1):
            ai_str = f"{item['ai'] * 100:.0f}%" if isinstance(item["ai"], (int, float)) else "—"
            rows.append([i, item["name"], f"{item['score']:.2f}", ai_str, item["priority"]])
        add_table(doc, ["排名", "开发者", "综合评分", "AI 采纳率", "改进优先级"], rows)

    if len(scored) > 10:
        add_paragraph(doc, "### 3.2 Bottom 5 低分开发者", bold=True, size=12)
        rows = []
        for i, item in enumerate(scored[-5:], len(scored) - 4):
            ai_str = f"{item['ai'] * 100:.0f}%" if isinstance(item["ai"], (int, float)) else "—"
            rows.append([i, item["name"], f"{item['score']:.2f}", ai_str, item["priority"]])
        add_table(doc, ["排名", "开发者", "综合评分", "AI 采纳率", "改进优先级"], rows)

    if not scored:
        add_paragraph(doc, "本周无有效评分数据。")


def chapter_improvements(doc, overview):
    add_heading(doc, "第四章 改进建议汇总", level=1)
    devs = overview["developers"]
    buckets = {
        "4.1 文档结构方面": "document_structure",
        "4.2 内容完整性方面": "content_completeness",
        "4.3 反模式修正方面": "anti_pattern_correction",
        "4.4 AI 采纳优化方面": "ai_adoption_optimization",
        "4.5 AI 使用优化方面": "ai_usage_optimization",
    }
    for label, key in buckets.items():
        all_items = []
        for d in devs:
            sug = safe_get(d, "improvement_suggestions", default={}) or {}
            items = sug.get(key) or []
            if isinstance(items, list):
                for it in items:
                    if it and str(it).strip():
                        all_items.append(str(it).strip())
            elif items:
                all_items.append(str(items))
        if not all_items:
            continue
        add_heading(doc, label, level=2)
        counter = Counter(all_items)
        top = counter.most_common(8)
        rows = [[s, str(c)] for s, c in top]
        add_table(doc, ["建议", "出现次数"], rows)


def chapter_gaps(doc, overview):
    add_heading(doc, "第五章 不足与改进建议", level=1)
    devs = overview["developers"]
    no_data = [safe_get(d, "name", default="未知") for d in devs if not safe_get(d, "overall_score")]
    no_doc = [safe_get(d, "name", default="未知") for d in devs if safe_get(d, "ai_adoption_rate", "data_combination_type") == "code_only"]
    no_code = [safe_get(d, "name", default="未知") for d in devs if safe_get(d, "ai_adoption_rate", "data_combination_type") == "doc_only"]
    anti = [
        safe_get(d, "name", default="未知")
        for d in devs
        if safe_get(d, "anti_patterns", "has_issues") is True
    ]

    add_paragraph(doc, "### 5.1 数据缺失问题分析", bold=True, size=12)
    add_paragraph(
        doc,
        f"本次评审发现 {len(no_data)} 位开发者存在数据缺失问题（已按 0 分计入总评），主要表现为：",
    )
    rows = [
        ["未提交（已计 0 分）", f"{len(no_data)} 人", ", ".join(no_data[:10]) + ("..." if len(no_data) > 10 else "")],
        ["缺文档（仅有代码）", f"{len(no_doc)} 人", ", ".join(no_doc[:10]) or "—"],
        ["缺代码（仅有文档）", f"{len(no_code)} 人", ", ".join(no_code[:10]) or "—"],
        ["反模式命中", f"{len(anti)} 人", ", ".join(anti[:10]) or "—"],
    ]
    add_table(doc, ["问题类型", "人数", "涉及人员（部分）"], rows)

    add_paragraph(doc, "### 5.2 系统性解决方案", bold=True, size=12)
    for s in [
        "建立文档驱动的开发流程：要求每次开发任务必须先完成文档编写，再进行代码实现",
        "统一提交规范：代码提交记录应包含文档编号或关联信息，便于追溯",
        "强化审核机制：对 AI 生成内容进行标注和二次审核，增加个性化内容",
        "建立知识沉淀机制：记录开发过程中的问题解决经验和技术决策，形成团队知识库",
    ]:
        add_bullet(doc, s)

    add_paragraph(doc, "### 5.3 文档驱动开发流程优化建议", bold=True, size=12)
    for s in [
        "采用三层文档架构：规格说明（SPEC）→ 任务清单（TASKS）→ 验收检查表（CHECKLIST）",
        "保持文档间引用关系：SPEC§N → TASKS§N → CHECKLIST§N",
        "记录 AI 辅助使用：在文档中标注 AI 生成内容，记录 AI 辅助的具体场景和效果",
        "定期回顾更新：确保文档与代码同步演进，避免文档与实现脱节",
    ]:
        add_bullet(doc, s)


def chapter_missing_inline(doc, not_submitted, not_reviewed, total_members, actually_attended, reviewed_count):
    """嵌入第五章 5.x，拆分两类漏检（V2.2）"""
    add_heading(doc, "5.x 漏检汇总", level=2)
    add_paragraph(
        doc,
        f"应到 {total_members} 人，周目录实到 {actually_attended} 人，已评审 {reviewed_count} 人。",
    )
    rows = [
        ["真未提交 (A)", f"{len(not_submitted)} 人", "周目录无子目录，已按 0 分计入总评"],
        ["漏评审 (B)", f"{len(not_reviewed)} 人", "周目录有材料但模型未评审，需补件重评"],
    ]
    add_table(doc, ["漏检类型", "人数", "处理方式"], rows)

    if not_submitted:
        add_paragraph(doc, "5.x.1 真未提交者名单（已计 0 分）", bold=True, size=11)
        rows = [[i + 1, name] for i, name in enumerate(not_submitted)]
        add_table(doc, ["序号", "姓名"], rows)
    if not_reviewed:
        add_paragraph(doc, "5.x.2 漏评审名单（已补件重评）", bold=True, size=11)
        rows = [[i + 1, name] for i, name in enumerate(not_reviewed)]
        add_table(doc, ["序号", "姓名"], rows)
    if not not_submitted and not not_reviewed:
        add_paragraph(doc, "本周无漏检，全员已评审。")


def chapter_full_compare(doc, trend_md):
    """第六章 整体趋势变化分析"""
    add_heading(doc, "第六章 整体趋势变化分析", level=1)
    add_paragraph(
        doc,
        "本章将本周数据与历史周数据进行对比，展示整体趋势与个体变化。",
    )
    if not trend_md or not trend_md.exists():
        add_paragraph(doc, "无趋势分析文件，跳过趋势渲染。")
        return
    text = trend_md.read_text(encoding="utf-8")
    lines = text.splitlines()
    i = 0
    while i < len(lines):
        ln = lines[i]
        if ln.startswith("## "):
            heading_text = ln.lstrip("# ").strip()
            add_heading(doc, heading_text, level=2)
            i += 1
            continue
        if ln.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|") and set(lines[i + 1].replace("|", "").strip()) <= {"-", " "}:
            header = [c.strip() for c in ln.strip("|").split("|")]
            i += 2
            data_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                row = [c.strip() for c in lines[i].strip("|").split("|")]
                data_rows.append(row)
                i += 1
            if data_rows:
                add_table(doc, header, data_rows)
            continue
        if ln.startswith(">"):
            add_paragraph(doc, ln.lstrip("> ").strip())
        elif ln.strip().startswith("- "):
            add_bullet(doc, ln.strip()[2:])
        elif ln.strip():
            add_paragraph(doc, ln.strip())
        i += 1


def chapter_conclusion(doc, overview):
    """总结（独立章节，参考样例 AI辅助编程评审报告0511-0517.docx）"""
    add_heading(doc, "总结", level=1)
    devs = overview["developers"]
    scored = [d for d in devs if isinstance(safe_get(d, "overall_score"), (int, float))]
    valid = len(scored)
    if valid:
        avg_score = sum(safe_get(d, "overall_score") for d in scored) / valid
        add_paragraph(
            doc,
            f"通过本次 AI 辅助编程评审，可以看出团队在文档驱动开发方面持续优化。本周共评审 {valid} 位有效数据开发者，平均综合评分 {avg_score:.2f}。",
        )
    else:
        add_paragraph(
            doc,
            "通过本次 AI 辅助编程评审，团队在文档驱动开发方面持续优化。",
        )
    add_paragraph(doc, "未来改进重点方向：")
    for s in [
        "继续强化文档与代码的一致性，确保文档真实反映开发过程",
        "深化 AI 辅助的使用层次，从文档生成扩展到方案优化、技术审查等更广泛领域",
        "建立知识沉淀机制，将开发经验系统化记录，形成团队资产",
        "关注数据缺失的开发者，提供必要的培训和指导",
        "通过持续改进，团队将能够更好地利用 AI 辅助编程工具，提升整体开发效率",
    ]:
        add_bullet(doc, s)


def chapter_baseline_compare(doc, db_path, current_week, current_names):
    """第七章 全量比对（最近 4 次），从 SQLite 读取"""
    add_heading(doc, "第七章 全量比对（最近 4 次）", level=1)
    add_paragraph(
        doc,
        f"本章按人列出最近 4 次的总分（数据源: {db_path.name}），含本周提交的所有开发者。",
    )

    if not db_path.exists():
        add_paragraph(doc, f"未找到 baseline 数据库: {db_path}")
        return

    conn = sqlite3.connect(str(db_path))
    try:
        cur = conn.cursor()
        rows_data = []
        for name in current_names:
            cur.execute(
                """
                SELECT week, overall_score
                FROM review_history
                WHERE name = ?
                ORDER BY week DESC
                LIMIT 4
                """,
                (name,),
            )
            history = cur.fetchall()
            scores = [s for _, s in history]
            scores.reverse()
            while len(scores) < 4:
                scores.insert(0, None)
            first, last = (scores[0], scores[-1]) if scores and scores[0] is not None else (None, None)
            if first is None and last is None:
                trend = TREND_NEW
            elif first is None or last is None:
                trend = TREND_NEW
            elif last > first + 0.5:
                trend = TREND_UP
            elif last < first - 0.5:
                trend = TREND_DOWN
            else:
                trend = TREND_FLAT
            valid = [s for s in scores if s is not None]
            avg = sum(valid) / len(valid) if valid else None
            rows_data.append({
                "name": name,
                "scores": scores,
                "trend": trend,
                "avg": avg,
                "this_week": last,
            })

        rows_data.sort(key=lambda r: (r["this_week"] is None, -(r["this_week"] or 0)))

        rows = []
        for r in rows_data:
            score_cells = [f"{s:.2f}" if s is not None else "—" for s in r["scores"]]
            rows.append([
                r["name"],
                *score_cells,
                r["trend"],
                f"{r['avg']:.2f}" if r["avg"] is not None else "—",
            ])
        add_table(
            doc,
            ["开发者", "第1次(最早)", "第2次", "第3次", "第4次(最新)", "趋势", "平均"],
            rows,
        )

        cur.execute(
            """
            SELECT COUNT(DISTINCT name), COUNT(DISTINCT week)
            FROM review_history
            """
        )
        total_devs, total_weeks = cur.fetchone()
        add_paragraph(
            doc,
            f"备注: 基准库共记录 {total_devs} 位开发者、{total_weeks} 个周次的评审数据。",
        )
    finally:
        conn.close()


# ========================================================================
# 主流程
# ========================================================================

def main():
    parser = argparse.ArgumentParser(description="生成 AI 辅助编程周报 DOCX")
    parser.add_argument("--json", required=True, help="review_results_*.json 路径")
    parser.add_argument("--excel", required=True, help="文档审查结果_*.xlsx 路径")
    parser.add_argument("--trend-md", default=None, help="趋势分析_*.md 路径")
    parser.add_argument("--missing-json", default=None, help="漏检结果 JSON 路径")
    parser.add_argument("--db", required=True, help="baseline.db 路径（第七章用）")
    parser.add_argument("--output", required=True, help="输出 docx 路径")
    parser.add_argument("--week", required=True, help="周次标签，如 0525-0531")
    args = parser.parse_args()

    json_path = Path(args.json)
    excel_path = Path(args.excel)
    trend_md = Path(args.trend_md) if args.trend_md else None
    missing_json = Path(args.missing_json) if args.missing_json else None
    db_path = Path(args.db)
    output_path = Path(args.output)
    week = args.week

    overview = collect_overview(json_path, week)
    excel_headers, excel_rows = load_excel_rows(excel_path)
    _ = (excel_headers, excel_rows)

    missing = []
    not_submitted = []
    not_reviewed = []
    total_members = 0
    actually_attended = 0
    reviewed_count = overview["count"]
    if missing_json and missing_json.exists():
        m = load_json(missing_json)
        not_submitted = m.get("not_submitted", [])
        not_reviewed = m.get("not_reviewed", [])
        if not not_submitted and not not_reviewed:
            not_submitted = m.get("missing", [])
        missing = sorted(set(not_submitted) | set(not_reviewed))
        total_members = m.get("total_members", 0)
        actually_attended = m.get("actually_attended", 0)
        reviewed_count = m.get("reviewed_count", reviewed_count)

    current_names = []
    seen = set()
    for d in overview["developers"]:
        n = safe_get(d, "name")
        if n and n not in seen:
            current_names.append(str(n).strip())
            seen.add(n)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), "微软雅黑")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")

    title = doc.add_heading("", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("AI 辅助编程评审报告")
    tr.font.color.rgb = BLACK
    tr.font.size = Pt(22)
    tr.bold = True
    set_run_fonts(tr)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("开发者评审内容汇总与分析")
    sr.font.size = Pt(14)
    sr.bold = True
    sr.font.color.rgb = BLACK
    set_run_fonts(sr)

    add_paragraph(
        doc,
        f"周次: {week}    生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        size=10,
    )

    chapter_summary(doc, overview, json_path)
    chapter_stats(doc, overview)
    chapter_top_bottom(doc, overview)
    chapter_improvements(doc, overview)
    chapter_gaps(doc, overview)
    chapter_missing_inline(doc, not_submitted, not_reviewed, total_members, actually_attended, reviewed_count)
    chapter_full_compare(doc, trend_md)
    chapter_conclusion(doc, overview)
    chapter_baseline_compare(doc, db_path, week, current_names)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"报告已生成: {output_path}")
    print(f"  章节: 7 + 1（标题/概述/统计/Top/改进/不足/5.x漏检/趋势/总结/全量比对）")
    print(f"  Baseline: {db_path}")
    print(f"  本周提交者: {len(current_names)} 人")


if __name__ == "__main__":
    main()
