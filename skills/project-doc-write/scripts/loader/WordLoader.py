#!/usr/bin/python
# -*- coding:utf-8 -*-
"""
WordLoader 模块（project-doc-write skill 自带）

忠实复制自 agent-user-mangerment（与 query 版同构，独立副本）。

Date: 2026-06-10
Author: project-doc skill 套件
"""

from pathlib import Path
from typing import Optional, List, Dict, Any

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None  # type: ignore

try:
    from langchain_core.documents import Document  # type: ignore
except Exception:
    Document = None  # type: ignore


class _SimpleDoc:
    __slots__ = ("page_content", "metadata")

    def __init__(self, page_content: str = "", metadata=None):
        self.page_content = page_content
        self.metadata = metadata or {}


class WordLoader:
    LOADER_NAME = "WordLoader"

    def __init__(self, file_path: str, load_method: str = "default", **kwargs):
        self.file_path = Path(file_path)
        self.load_method = load_method
        self._doc = None

    def _get_document(self):
        if DocxDocument is None:
            raise ImportError("缺少依赖 python-docx，请先安装：pip install python-docx")
        if self._doc is None:
            if not self.file_path.exists():
                raise FileNotFoundError(f"文件不存在: {self.file_path}")
            self._doc = DocxDocument(str(self.file_path))
        return self._doc

    def _make_doc(self, page_content: str, metadata: Dict[str, Any]):
        DocCls = Document if Document is not None else _SimpleDoc
        return DocCls(page_content=page_content, metadata=metadata)

    def load(self) -> List[Any]:
        doc = self._get_document()
        documents: List[Any] = []
        for idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                documents.append(self._make_doc(
                    page_content=para.text,
                    metadata={"source": str(self.file_path), "type": "paragraph", "index": idx},
                ))
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                row_text = [cell.text for cell in row.cells]
                cell_content = " | ".join(row_text)
                if cell_content.strip():
                    documents.append(self._make_doc(
                        page_content=cell_content,
                        metadata={
                            "source": str(self.file_path),
                            "type": "table",
                            "table_index": table_idx,
                            "row_index": row_idx,
                        },
                    ))
        return documents

    def lazy_load(self) -> List[Any]:
        return self.load()

    @property
    def exists(self) -> bool:
        return self.file_path.exists()


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("用法: python WordLoader.py <docx_path>")
        sys.exit(1)
    loader = WordLoader(sys.argv[1])
    if loader.exists:
        docs = loader.load()
        print(f"加载 {len(docs)} 个片段")
    else:
        print(f"文件不存在: {loader.file_path}")
