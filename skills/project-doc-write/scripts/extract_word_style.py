#!/usr/bin/python
# -*- coding:utf-8 -*-
"""
extract_word_style.py

从项目内 docx（如 03_技术文档及评审/01_实施方案/实施方案V1.0.docx）抽取
样式基准：封面、字体、字号、行间距、页眉页脚、表格样式、段落间距等。

输出到 stdout（JSON），可重定向到文件作为缓存。
本脚本是 write skill 内部的"一次性"工具：首次跑通后，样式信息
会被固化为 references/word_格式范本_规则.md 的写死内容。

Date: 2026-06-11
Author: project-doc-write skill
"""

import sys
from pathlib import Path

if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8")

import argparse
import json
from collections import OrderedDict

_SKILL_SCRIPTS = Path(__file__).resolve().parent
sys.path.insert(0, str(_SKILL_SCRIPTS))


def extract_styles(docx_path: Path) -> dict:
    """从 docx 抽取样式基准。"""
    try:
        from docx import Document
    except ImportError:
        raise SystemExit("缺少依赖 python-docx，请先安装：pip install python-docx")

    if not docx_path.exists():
        raise FileNotFoundError(f"文件不存在: {docx_path}")

    doc = Document(str(docx_path))
    styles_data = OrderedDict()

    for p in doc.paragraphs:
        if not p.text.strip():
            continue
        style_name = (p.style.name or "").strip() if p.style else ""
        try:
            font = p.font
        except (AttributeError, KeyError):
            font = None
        pf = p.paragraph_format
        try:
            font_name = font.name if font and font.name else ""
        except (AttributeError, KeyError, TypeError):
            font_name = ""
        try:
            font_size_pt = float(font.size.pt) if font and font.size else None
        except (AttributeError, KeyError, TypeError, ValueError):
            font_size_pt = None
        try:
            bold = font.bold if font and font.bold is not None else None
        except (AttributeError, KeyError, TypeError):
            bold = None
        try:
            alignment = str(p.alignment) if p.alignment is not None else None
        except (AttributeError, KeyError, TypeError):
            alignment = None
        try:
            space_before_pt = float(pf.space_before.pt) if pf.space_before else None
        except (AttributeError, KeyError, TypeError, ValueError):
            space_before_pt = None
        try:
            space_after_pt = float(pf.space_after.pt) if pf.space_after else None
        except (AttributeError, KeyError, TypeError, ValueError):
            space_after_pt = None
        try:
            line_spacing = pf.line_spacing
        except (AttributeError, KeyError, TypeError):
            line_spacing = None
        try:
            line_spacing_rule = str(pf.line_spacing_rule) if pf.line_spacing_rule else None
        except (AttributeError, KeyError, TypeError):
            line_spacing_rule = None
        entry = {
            "text_preview": p.text.strip()[:60],
            "style_name": style_name,
            "font_name": font_name,
            "font_size_pt": font_size_pt,
            "bold": bold,
            "alignment": alignment,
            "space_before_pt": space_before_pt,
            "space_after_pt": space_after_pt,
            "line_spacing": line_spacing,
            "line_spacing_rule": line_spacing_rule,
        }
        key = style_name or "Normal"
        styles_data.setdefault(key, []).append(entry)

    # 表格样式抽样
    tables_data = []
    for ti, table in enumerate(doc.tables[:3]):
        rows = len(table.rows)
        cols = len(table.columns) if table.columns else 0
        tbl_style = (table.style.name if table.style else "None")
        tables_data.append({
            "index": ti,
            "rows": rows,
            "cols": cols,
            "style_name": tbl_style,
        })

    # 节信息（页眉页脚/页面尺寸）
    sections_data = []
    for si, section in enumerate(doc.sections):
        sections_data.append({
            "index": si,
            "page_height_cm": round(section.page_height.cm, 2) if section.page_height else None,
            "page_width_cm": round(section.page_width.cm, 2) if section.page_width else None,
            "top_margin_cm": round(section.top_margin.cm, 2) if section.top_margin else None,
            "bottom_margin_cm": round(section.bottom_margin.cm, 2) if section.bottom_margin else None,
            "left_margin_cm": round(section.left_margin.cm, 2) if section.left_margin else None,
            "right_margin_cm": round(section.right_margin.cm, 2) if section.right_margin else None,
        })

    return {
        "source": str(docx_path),
        "styles": styles_data,
        "tables": tables_data,
        "sections": sections_data,
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
    }


def main() -> int:
    parser = argparse.ArgumentParser(
        description="从项目内 docx 抽取样式基准",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "示例:\n"
            "  python extract_word_style.py --docx \"D:\\\\项目文档\\\\...\\\\实施方案.docx\" --output style_benchmark.json\n"
        ),
    )
    parser.add_argument("--docx", required=True, help="项目内 docx 格式范本路径")
    parser.add_argument("--output", default=None, help="输出 JSON 路径（默认 stdout）")
    args = parser.parse_args()

    try:
        benchmark = extract_styles(Path(args.docx))
    except FileNotFoundError as e:
        print(f"[ERR] {e}", file=sys.stderr)
        return 1
    except Exception as e:
        print(f"[ERR] 抽取失败: {e}", file=sys.stderr)
        return 3

    text = json.dumps(benchmark, ensure_ascii=False, indent=2)
    if args.output:
        out_path = Path(args.output)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(text, encoding="utf-8")
        print(f"[OK] 样式基准已抽取 → {out_path}")
        print(f"     段落样式: {len(benchmark['styles'])}, 表格: {benchmark['table_count']}, 节: {len(benchmark['sections'])}")
    else:
        print(text, flush=True)

    return 0


if __name__ == "__main__":
    sys.exit(main())
